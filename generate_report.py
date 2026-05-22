import asyncio
from playwright.async_api import async_playwright
from docx import Document
from docx.shared import Inches
import os
import sys

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

PAGES = [
    ("Главная страница (Дашборд)", "/", "Главная страница агрегирует ключевую информацию по текущему состоянию системы."),
    ("Справочник единиц измерения", "/units/", "Управление единицами измерения (шт, кг, метры и т.д.)."),
    ("Справочник должностей", "/positions/", "Управление списком должностей сотрудников склада."),
    ("Справочник сотрудников", "/employees/", "Модуль для работы с кадрами."),
    ("Справочник сырья", "/raw_materials/", "Справочник сырьевых материалов, используемых в производстве."),
    ("Справочник готовой продукции", "/finished_products/", "Учёт произведенных товаров (одежды)."),
    ("Состав продукции (Ингредиенты)", "/ingredients/", "Модуль для настройки спецификаций (рецептур) производства одежды."),
    ("Закупка сырья (Форма)", "/purchase/add/", "Форма для оформления поступления сырья на склад."),
    ("Производство", "/production/", "Процесс создания готовой продукции из сырья."),
    ("Оформление продажи готовой продукции", "/sale/add/", "Оформление отгрузки готовой продукции клиентам."),
    ("Состояние бюджета", "/budget/", "Отображение текущего финансового состояния предприятия."),
    ("Выплата зарплаты", "/salaries/", "Оформление выплат сотрудникам согласно их окладам."),
    ("Бизнес-кредиты", "/loans/", "Форма для привлечения заёмных средств."),
    ("Аналитический дашборд", "/analytics/dashboard/", "Графическое представление ключевых метрик.")
]

async def capture_screenshots():
    BASE_URL = "http://127.0.0.1:8080"
    os.makedirs("screenshots", exist_ok=True)
    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page(viewport={"width": 1280, "height": 800})
        
        for name, path, _ in PAGES:
            try:
                print(f"Loading {name}...")
                await page.goto(f"{BASE_URL}{path}", wait_until="networkidle")
                filename = f"screenshots/{name}.png"
                await page.screenshot(path=filename)
                print(f"Captured: {name}")
            except Exception as e:
                print(f"Failed to capture {name}: {e}")
        
        await browser.close()

def create_docx():
    doc = Document()
    doc.add_heading('Отчет по проекту: Система управления складом одежды', 0)
    doc.add_paragraph('Автоматически сгенерированный отчёт со скриншотами основных функций системы.')
    
    for name, _, desc in PAGES:
        doc.add_heading(name, level=1)
        doc.add_paragraph(desc)
        filename = f"screenshots/{name}.png"
        if os.path.exists(filename):
            try:
                doc.add_picture(filename, width=Inches(6.0))
            except Exception as e:
                doc.add_paragraph(f"[Ошибка вставки изображения: {e}]")
        else:
            doc.add_paragraph("[Скриншот не найден или страница не загрузилась]")
            
    doc.save("project_report.docx")
    print("Docx created and saved as project_report.docx")

if __name__ == "__main__":
    asyncio.run(capture_screenshots())
    create_docx()
