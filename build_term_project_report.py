"""Generate a DBMS term project report in DOCX matching the sample layout.

Sample reference: DBMS_TermProject_Kanybekova_SEengl-1-23 (2).pdf
Adapted to: WarehouseOfClothes (Django + MSSQL, garment factory).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BASE = os.path.dirname(os.path.abspath(__file__))
SHOTS = os.path.join(BASE, "screenshots")

doc = Document()

# ── global styles ──
for s in doc.sections:
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(1.5)
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(4)

# helpers ----------------------------------------------------------
def p(text="", bold=False, italic=False, size=12, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
      indent=True, color=None):
    para = doc.add_paragraph()
    para.paragraph_format.alignment = align
    if indent:
        para.paragraph_format.first_line_indent = Cm(1.0)
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return para

def h1(text):
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)
    r = para.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)
    return para

def h2(text):
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    r = para.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    return para

def h3(text):
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)
    r = para.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    return para

def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Cm(1.0)
    para.paragraph_format.space_after = Pt(0)
    r = para.runs[0] if para.runs else para.add_run("")
    para.add_run(text).font.name = "Times New Roman"
    para.runs[-1].font.size = Pt(12)
    return para

def caption(text):
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_after = Pt(10)
    r = para.add_run(text)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

def picture(path, width_inches=5.5, cap=None):
    if not os.path.exists(path):
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.add_run(f"[image missing: {os.path.basename(path)}]").italic = True
        if cap:
            caption(cap)
        return
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.add_run().add_picture(path, width=Inches(width_inches))
    if cap:
        caption(cap)

def code_block(text):
    """Mono-spaced code paragraph with light gray fill."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(6)
    r = para.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    # background shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)

def shade_cell(cell, color="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None, header_color="C5677B"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, t in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = para.add_run(t)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], header_color)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            r = para.add_run(str(val))
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ============================================================
# TITLE PAGE
# ============================================================
def center_line(txt, bold=False, size=12, space_before=0, space_after=0):
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.first_line_indent = 0
    r = para.add_run(txt)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return para

center_line("MINISTRY OF SCIENCE, HIGHER EDUCATION AND INNOVATIONS", size=12, space_before=24)
center_line("OF THE KYRGYZ REPUBLIC", size=12, space_after=18)
center_line("KYRGYZ STATE TECHNICAL UNIVERSITY named after I. Razzakov", size=12, space_after=18)
center_line("INSTITUTE OF INFORMATION TECHNOLOGIES", size=12)
center_line('DEPARTMENT OF "SOFTWARE OF COMPUTER SYSTEMS"', size=12, space_after=18)
center_line("710400 “Software Engineering”", bold=True, size=13, space_after=80)

center_line("Term project", bold=True, size=28, space_after=24)

center_line("For the course: Database Management Systems", italic := False, size=13, space_after=8)
center_line("On the topic: Development of a web application for managing", size=13)
center_line("a clothing factory warehouse (WarehouseOfClothes)", size=13, space_after=140)

# author / instructor blocks (right aligned)
def right_line(txt, bold=False, size=12, space_after=0):
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.first_line_indent = 0
    r = para.add_run(txt)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)

right_line("Prepared by:", bold=True, size=12)
right_line("Student of group SE(b)-1-23", size=12)
right_line("Klyuchevsky D.", bold=True, size=12, space_after=18)
right_line("Instructor:", bold=True, size=12)
right_line("Senior Lecturer", size=12)
right_line("Makieva Z.D.", size=12, space_after=120)

center_line("Bishkek, 2026", size=12)

doc.add_page_break()


# ============================================================
# TABLE OF CONTENTS
# ============================================================
h1("Table of Content")

toc = [
    ("1. Technical Specification", "2"),
    ("    1.1. Introduction", "2"),
    ("    1.2. Objective of the Work", "2"),
    ("    1.3. Problem Specification", "3"),
    ("    1.4. Purpose and Goals of System Creation", "3"),
    ("    1.5. Functional Requirements", "4"),
    ("    1.6. Client-Side Requirements", "4"),
    ("2. Server Side", "5"),
    ("    2.1. Database", "5"),
    ("    2.2. Tables", "6"),
    ("    2.3. Stored procedures", "13"),
    ("    2.4. Views", "15"),
    ("    2.5. Triggers", "17"),
    ("3. Client side", "19"),
    ("    3.1. Description of project creation", "19"),
    ("    3.2. More detailed analysis", "22"),
    ("        1. Models (Entity)", "22"),
    ("        2. Views (Controllers)", "24"),
    ("        3. Forms", "26"),
    ("        4. URL routing", "27"),
    ("        5. Permissions and Roles", "28"),
    ("        6. Templates (Presentation)", "29"),
    ("4. Launching the program", "32"),
    ("5. Screenshots of the program", "33"),
    ("6. Conclusion", "42"),
    ("7. List of Used Literature", "42"),
    ("8. Application appendix", "43"),
]
for item, page in toc:
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = 0
    para.paragraph_format.space_after = Pt(2)
    tab = OxmlElement("w:tabs")
    t = OxmlElement("w:tab")
    t.set(qn("w:val"), "right")
    t.set(qn("w:leader"), "dot")
    t.set(qn("w:pos"), "9000")
    tab.append(t)
    para._p.get_or_add_pPr().append(tab)
    r = para.add_run(item)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    if item.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.")) and not item.startswith("    "):
        r.bold = True
    para.add_run("\t" + page).font.size = Pt(12)

doc.add_page_break()


# ============================================================
# 1. TECHNICAL SPECIFICATION
# ============================================================
h1("1. Technical Specification")

h2("1.1. Introduction")
p("In the context of the modern garment-industry market, effective management of "
  "warehouse, production and financial resources is becoming a key success factor "
  "for a clothing factory. Such an enterprise needs to control the receipt and "
  "consumption of raw materials (fabric, threads, accessories), the manufacture of "
  "finished items, sales to clients, salary payouts to employees and the use of "
  "borrowed funds — all on a single budget.")
p("The web application WarehouseOfClothes was developed in order to replace manual "
  "(paper-based and spreadsheet) accounting with a unified digital management system. "
  "Without such a system, the factory cannot reliably track raw-material stock, "
  "control production cost, manage finished-product inventory, prevent sales below "
  "cost, or analyse who and when performed each operation. Implementing this web "
  "application is therefore a relevant task that allows the factory to operate "
  "more efficiently and to make data-driven decisions.")

h2("1.2. Objective of the Work")
p("The objective of the work is to design and implement a comprehensive web "
  "application for a clothing factory warehouse. The main task is to create tools "
  "that will allow authorized users to:")
bullet("view and manage the factory budget;")
bullet("record and pay off business loans;")
bullet("manage employees, positions and salary payments;")
bullet("track purchases of raw materials with units of measurement;")
bullet("manage the catalogue of finished products and their recipes (ingredients);")
bullet("register production operations with raw-material consumption check;")
bullet("register sales of finished products with automatic price calculation;")
bullet("automatically update the budget on every financial operation;")
bullet("submit and process production requests with multi-stage workflow;")
bullet("view analytics, top products and the ratio of revenue to expenses.")

h2("1.3. Problem Specification")
p("During the analysis of the existing management processes of a clothing factory, "
  "the following problems were identified:")
bullet("Absence of a unified accounting tool — purchases, production and sales are "
       "kept in separate Excel sheets, making it difficult to track real stock.")
bullet("Lack of automatic budget control: any cash operation is registered manually, "
       "which produces arithmetic errors and disagreements between departments.")
bullet("Inability to forbid a sale below cost or production with insufficient raw "
       "materials — such checks have to be done by the operator, who can forget them.")
bullet("Absence of role-based access control: every user can see and modify any data, "
       "which contradicts internal company policy.")
bullet("No analytical reports — the manager has no consolidated view of top-selling "
       "products, total expenses and current stock.")

h2("1.4. Purpose and Goals of System Creation")
h3("1.4.1. Purpose")
p("Development of a web application for managing a clothing-factory warehouse "
  "to provide directors, accountants, warehouse keepers, production specialists and "
  "sales managers with effective tools for comprehensive resource management, "
  "simplifying the processes of budget control, procurement, production, "
  "sales and HR.")
h3("1.4.2. Goals")
p("The main goals of creating the new system are:")
bullet("simplify the process of warehouse accounting and financial planning;")
bullet("ensure quick access to information about loans and their repayment status;")
bullet("automate employee management, salary payments and position tracking;")
bullet("provide transparent control over raw-material purchases and units of measure;")
bullet("enable efficient management of the finished-product line, production and sales;")
bullet("implement role-based access so that each user only sees their working area.")

h2("1.5. Functional Requirements")
bullet("Creating and editing the budget record.")
bullet("Recording business loans (up to three simultaneously) and paying them off.")
bullet("Creating, editing and viewing employees, positions and units of measurement.")
bullet("Processing salary payments with automatic budget deduction.")
bullet("Recording raw-material purchases through a stored procedure that verifies "
       "the budget and atomically updates stock and budget via an INSERT trigger.")
bullet("Managing the catalogue of finished products and their recipes (ingredients).")
bullet("Performing a production operation only if raw materials are sufficient; "
       "the system automatically consumes raw materials and increases the product stock.")
bullet("Recording sales with auto-calculated price (cost × 1.30); the budget is "
       "increased on every sale.")
bullet("Production requests: an end-to-end workflow (created → checking raw stock → "
       "purchasing raw materials → producing → selling → completed).")
bullet("Three SQL views give read-only reports: raw-stock view, finished-product-stock "
       "view, and materials-needed-for-production view.")
bullet("Analytics dashboard with charts of top sold products and revenue vs expenses.")
bullet("Role-based access control: director (full access), accountant (budget, salary, "
       "loans), warehouse keeper (purchases), production worker (production, recipes), "
       "sales manager (sales).")
bullet("Account management: the director can create and revoke accounts of employees.")

h2("1.6. Client-Side Requirements")
bullet("The client can use one of the following web browsers: Google Chrome, Mozilla "
       "Firefox, Microsoft Edge, Yandex Browser.")
bullet("To use the system, the user must authenticate with a valid username and password.")
bullet("After logging in, access to specific functions (budget editing, loan management, "
       "purchases, sales, salary payouts) is controlled by role-based permissions.")
bullet("The interface automatically adapts to the user’s role, hiding irrelevant "
       "sections and menu items.")
bullet("The user interface must be simple, clean and convenient for daily work.")
bullet("The main pages of the application are accessible through the top navigation bar.")
bullet("The system must display important information — current budget, active loans, "
       "the authorized user — in a visible area of the interface.")
bullet("Pages must support basic operations: adding, editing, deleting and viewing "
       "history of records.")
bullet("Error messages must be displayed in a user-friendly way when a business rule "
       "is violated, e.g. insufficient budget, sale below cost, missing raw materials.")

doc.add_page_break()


# ============================================================
# 2. SERVER SIDE
# ============================================================
h1("2. Server Side")

h2("2.1. Database")
p("The application uses Microsoft SQL Server as its DBMS. The schema is generated "
  "and maintained by Django ORM migrations; all tables therefore receive the prefix "
  "warehouse_. Together with Django service tables (auth_user, django_session, "
  "django_migrations, etc.) the database contains 22 tables; 14 of them are the "
  "domain tables of the WarehouseOfClothes system. In addition the database has "
  "2 user-defined stored procedures, 1 INSERT trigger and 3 SQL views — all of "
  "them created in migration 0002_create_sql_logic.py.")

picture(os.path.join(SHOTS, "БД_Список_таблиц.png"), 5.5, "Image 1. The list of all tables of the system in MSSQL.")

picture(os.path.join(SHOTS, "БД_ER_диаграмма.png"), 6.2, "Image 1.5. Entity-Relationship diagram of the database.")


h2("2.2. Tables")

# Each domain table
tables_info = [
    ("Budget", "Stores the current value of the factory’s budget. There must "
     "always be exactly one row (an in-app guard in the model forbids creating a "
     "second instance).",
     [
        ("id", "BigAutoField", "Primary Key"),
        ("amount", "Float", "Current budget amount (KGS)."),
     ],
     "БД_Проект_Budget.png", "Image 2. Structure of the table Budget."),

    ("Unit", "Reference list of measurement units (pcs, kg, m, etc.) used both "
     "for raw materials and for finished products.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("name", "CharField(50), UNIQUE", "Name of the unit."),
     ],
     "БД_Проект_Unit.png", "Image 3. Structure of the table Unit."),

    ("Position", "Reference list of employee positions; each position is linked "
     "to one application role (director / accountant / warehouse / production / "
     "sales / none) that governs RBAC permissions.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("title", "CharField(100), UNIQUE", "Position title."),
        ("role", "CharField(20)", "Role-based-access role (choice field)."),
     ],
     "БД_Проект_Position.png", "Image 4. Structure of the table Position."),

    ("Employee", "Stores information about the factory employees.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("full_name", "CharField(200)", "Full name of the employee."),
        ("position_id", "FK → Position (PROTECT)", "Position of the employee."),
        ("salary", "Decimal(10,2)", "Monthly salary."),
        ("address", "CharField(255), nullable", "Residential address."),
        ("phone", "CharField(20), nullable", "Contact phone number."),
     ],
     "БД_Проект_Employee.png", "Image 5. Structure of the table Employee."),

    ("UserProfile", "Links a Django auth_user record with an Employee record and "
     "thus inherits the role from Employee.position.role. Used by the permission "
     "system.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("user_id", "OneToOne → auth_user (CASCADE)", "Django user."),
        ("employee_id", "OneToOne → Employee (PROTECT)", "Linked employee."),
     ],
     "БД_Проект_UserProfile.png", "Image 6. Structure of the table UserProfile."),

    ("RawMaterial", "Catalogue of raw materials with their current stock, "
     "average value of the stock and a fallback minimum price.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("name", "CharField(200)", "Name of the raw material."),
        ("unit_id", "FK → Unit (PROTECT)", "Unit of measurement."),
        ("quantity", "Float", "Current quantity on stock."),
        ("amount", "Float", "Total cost of the stock."),
        ("min_price", "Float", "Minimum / base price per 1 unit."),
     ],
     "БД_Проект_RawMaterial.png", "Image 7. Structure of the table RawMaterial."),

    ("FinishedProduct", "Catalogue of finished products with their stock and "
     "total stock value.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("name", "CharField(200)", "Name of the product."),
        ("unit_id", "FK → Unit (PROTECT)", "Unit of measurement."),
        ("quantity", "Float", "Quantity on stock."),
        ("amount", "Float", "Total cost of the stock."),
     ],
     "БД_Проект_FinishedProduct.png", "Image 8. Structure of the table FinishedProduct."),

    ("Ingredient", "Recipe table: which raw materials and in which proportion "
     "are required to manufacture one unit of a finished product. The pair "
     "(product, raw_material) is unique.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("product_id", "FK → FinishedProduct (CASCADE)", "Finished product."),
        ("raw_material_id", "FK → RawMaterial (PROTECT)", "Raw material."),
        ("quantity", "Float", "Raw material per 1 unit of product."),
     ],
     "БД_Проект_Ingredient.png", "Image 9. Structure of the table Ingredient."),

    ("RawPurchase", "Records every purchase of raw materials. The INSERT is "
     "performed via stored procedure sp_create_purchase; the AFTER-INSERT trigger "
     "trg_after_purchase updates stock and budget atomically.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("raw_material_id", "FK → RawMaterial (PROTECT)", "Purchased material."),
        ("quantity", "Float", "Quantity purchased."),
        ("amount", "Float", "Total purchase amount."),
        ("date", "DateTime auto_now_add", "Date of purchase."),
        ("employee_id", "FK → Employee (PROTECT)", "Employee who registered."),
     ],
     "БД_Проект_RawPurchase.png", "Image 10. Structure of the table RawPurchase."),

    ("Production", "Records every act of producing finished products.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("product_id", "FK → FinishedProduct (PROTECT)", "Manufactured product."),
        ("quantity", "Float", "Produced quantity."),
        ("date", "DateTime auto_now_add", "Date of production."),
        ("employee_id", "FK → Employee (PROTECT)", "Employee responsible."),
     ],
     "БД_Проект_Production.png", "Image 11. Structure of the table Production."),

    ("ProductSale", "Records sales of finished products. Sale price is "
     "calculated automatically as unit_cost × 1.30 × qty (the system forbids "
     "selling below cost).",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("product_id", "FK → FinishedProduct (PROTECT)", "Sold product."),
        ("quantity", "Float", "Sold quantity."),
        ("amount", "Float", "Total sale amount."),
        ("date", "DateTime auto_now_add", "Date of sale."),
        ("employee_id", "FK → Employee (PROTECT)", "Employee who registered."),
     ],
     "БД_Проект_ProductSale.png", "Image 12. Structure of the table ProductSale."),

    ("SalaryPayment", "Records salary payouts to employees. The corresponding "
     "amount is deducted from the budget by the view.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("employee_id", "FK → Employee (PROTECT)", "Receiving employee."),
        ("amount", "Float", "Amount paid."),
        ("date", "DateTime auto_now_add", "Payment date."),
     ],
     "БД_Проект_SalaryPayment.png", "Image 13. Structure of the table SalaryPayment."),

    ("BusinessLoan", "Records active loans of the factory. A business rule "
     "(enforced in the model) limits the number of simultaneously active loans "
     "to three.",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("amount", "Float", "Loan amount."),
        ("date_taken", "DateTime auto_now_add", "Date received."),
     ],
     "БД_Проект_BusinessLoan.png", "Image 14. Structure of the table BusinessLoan."),

    ("ProductionRequest", "Stores production requests filed by employees with "
     "a multi-step status workflow (created → checking raw stock → purchasing "
     "→ producing → selling → completed; or → error with a rejection reason).",
     [
        ("id", "BigAutoField", "Primary Key."),
        ("applicant_name", "CharField(200)", "Full name of the applicant."),
        ("product_id", "FK → FinishedProduct (CASCADE)", "Requested product."),
        ("quantity", "Float", "Requested quantity."),
        ("status", "CharField(50)", "Current status of the request."),
        ("rejection_reason", "TextField, nullable", "Reason if the request fails."),
        ("created_at", "DateTime auto_now_add", "Creation timestamp."),
        ("updated_at", "DateTime auto_now", "Last status change."),
     ],
     "БД_Проект_ProductionRequest.png", "Image 15. Structure of the table ProductionRequest."),
]

for n, (name, descr, fields, img, cap) in enumerate(tables_info, start=1):
    h3(f'{n}. “{name}” table')
    p(descr)
    p("Table structure:", bold=False, indent=False)
    add_table(["Column", "Type", "Description"], fields, col_widths=[4.5, 5.0, 7.0])
    picture(os.path.join(SHOTS, img), 4.5, cap)


# ─── 2.3 Stored procedures ─────────────────────────────────────
doc.add_page_break()
h2("2.3. Stored procedures")
picture(os.path.join(SHOTS, "БД_Хранимые_процедуры.png"), 4.5,
        "Image 16. List of user-defined stored procedures in DBMS.")

p("Two user-defined stored procedures are created in the database. Together they "
  "encapsulate the most critical write operation of the system — buying raw "
  "materials. The application calls only sp_create_purchase; it in turn checks "
  "the budget via sp_check_budget and, if the funds are sufficient, performs "
  "an INSERT into warehouse_rawpurchase. The actual update of stock and budget "
  "is then performed by the trigger trg_after_purchase (see § 2.5).")

add_table(["Name", "Description"],
          [
              ("sp_check_budget",
               "Checks whether the current budget covers the given amount. "
               "Reads the single row of warehouse_budget, returns 0 if there "
               "is enough money and 1 otherwise via an OUTPUT parameter."),
              ("sp_create_purchase",
               "Creates a new purchase: calls sp_check_budget, and only if the "
               "budget is sufficient INSERTs into warehouse_rawpurchase. The "
               "result (0 / 1) is returned via OUTPUT parameter so the Django "
               "view can show either a success message or an «insufficient "
               "funds» error."),
          ],
          col_widths=[4.5, 12.0])

h3("Code of sp_check_budget")
code_block(
"""CREATE OR ALTER PROCEDURE sp_check_budget
    @p_amount decimal(18,2),
    @result   int OUTPUT
AS
BEGIN
    DECLARE @v_budget decimal(18,2);
    SELECT TOP 1 @v_budget = amount
      FROM warehouse_budget
     ORDER BY id ASC;
    IF @v_budget IS NULL
        SET @result = 1;
    ELSE IF @v_budget >= @p_amount
        SET @result = 0;
    ELSE
        SET @result = 1;
END;"""
)

h3("Code of sp_create_purchase")
code_block(
"""CREATE OR ALTER PROCEDURE sp_create_purchase
    @p_raw_id      int,
    @p_qty         float,
    @p_amount      float,
    @p_date        datetime2,
    @p_employee_id int,
    @result        int OUTPUT
AS
BEGIN
    DECLARE @v_check int;
    EXEC sp_check_budget @p_amount, @v_check OUTPUT;

    IF @v_check = 0
    BEGIN
        INSERT INTO warehouse_rawpurchase
              (raw_material_id, quantity, amount, date, employee_id)
        VALUES(@p_raw_id, @p_qty, @p_amount, @p_date, @p_employee_id);
        SET @result = 0;
    END
    ELSE
    BEGIN
        SET @result = 1;
    END
END;"""
)

h3("Calling the procedure from the Django view")
code_block(
"""with connection.cursor() as cursor:
    cursor.execute(\"\"\"
        SET NOCOUNT ON;
        DECLARE @res INT;
        EXEC sp_create_purchase %s, %s, %s, %s, %s, @res OUTPUT;
        SELECT @res;
    \"\"\", [raw_id, quantity, amount, timezone.now(), employee_id])
    result = cursor.fetchone()[0]

if result == 0:
    messages.success(request, "Raw material purchased, budget updated.")
else:
    messages.error(request, "Insufficient budget for purchase.")"""
)


# ─── 2.4 Views (SQL) ───────────────────────────────────────────
doc.add_page_break()
h2("2.4. Views")
picture(os.path.join(SHOTS, "БД_Представления.png"), 4.5,
        "Image 17. List of SQL views in DBMS.")

p("Three SQL views are used as a thin read-only reporting layer. They are "
  "queried directly with raw SQL from Django views (see RawSQL examples in "
  "§ 3.2.2). Putting the join into a view keeps the Python code clean and "
  "lets the DBMS choose an optimal execution plan.")

add_table(["Name", "Description"],
          [
              ("v_raw_material_stock",
               "Joins warehouse_rawmaterial with warehouse_unit to expose "
               "current raw-material stock together with its unit name."),
              ("v_finished_product_stock",
               "Same idea for finished products: joins warehouse_finishedproduct "
               "with warehouse_unit."),
              ("v_materials_needed_for_production",
               "Joins warehouse_ingredient with warehouse_finishedproduct and "
               "warehouse_rawmaterial — gives a denormalised recipe view for "
               "the «materials needed» report."),
          ],
          col_widths=[5.5, 11.0])

h3("Code of v_raw_material_stock")
code_block(
"""CREATE OR ALTER VIEW v_raw_material_stock AS
SELECT
    r.id       AS raw_id,
    r.name     AS raw_name,
    u.name     AS unit_name,
    r.quantity,
    r.amount
FROM warehouse_rawmaterial r
JOIN warehouse_unit u ON r.unit_id = u.id;"""
)

h3("Code of v_finished_product_stock")
code_block(
"""CREATE OR ALTER VIEW v_finished_product_stock AS
SELECT
    p.id       AS product_id,
    p.name     AS product_name,
    u.name     AS unit_name,
    p.quantity,
    p.amount
FROM warehouse_finishedproduct p
JOIN warehouse_unit u ON p.unit_id = u.id;"""
)

h3("Code of v_materials_needed_for_production")
code_block(
"""CREATE OR ALTER VIEW v_materials_needed_for_production AS
SELECT
    p.id    AS product_id,
    p.name  AS product_name,
    r.id    AS raw_id,
    r.name  AS raw_name,
    i.quantity AS ingredient_qty
FROM warehouse_ingredient i
JOIN warehouse_finishedproduct p ON i.product_id      = p.id
JOIN warehouse_rawmaterial   r  ON i.raw_material_id  = r.id;"""
)


# ─── 2.5 Triggers ──────────────────────────────────────────────
doc.add_page_break()
h2("2.5. Triggers")
picture(os.path.join(SHOTS, "БД_Триггеры.png"), 4.5,
        "Image 18. List of triggers in DBMS.")

p("The system uses one INSERT trigger that keeps stock and budget in sync after "
  "every purchase of raw materials. Because the trigger is set-based it works "
  "correctly even for bulk INSERTs that aggregate by raw_material_id; the budget "
  "is decreased by SUM(amount), and stock is increased by the per-material totals.")

add_table(["Name", "Description"],
          [("trg_after_purchase",
            "Fires after a row is inserted into warehouse_rawpurchase. It deducts "
            "the total purchase amount from the single budget row and increases "
            "the stock (quantity, amount) of the affected raw materials."),
           ],
          col_widths=[5.0, 11.5])

h3("Code of trg_after_purchase")
code_block(
"""CREATE OR ALTER TRIGGER trg_after_purchase
ON warehouse_rawpurchase
AFTER INSERT
AS
BEGIN
    -- Deduct from Budget
    UPDATE warehouse_budget
       SET amount = amount - i.total_amount
      FROM warehouse_budget
      CROSS JOIN (
         SELECT SUM(amount) AS total_amount FROM inserted
      ) i
     WHERE id = (
         SELECT TOP 1 id FROM warehouse_budget ORDER BY id ASC
     );

    -- Update RawMaterial stock
    UPDATE r
       SET quantity = r.quantity + i.total_qty,
           amount   = r.amount   + i.total_amount
      FROM warehouse_rawmaterial r
      JOIN (
         SELECT raw_material_id,
                SUM(quantity) AS total_qty,
                SUM(amount)   AS total_amount
           FROM inserted
          GROUP BY raw_material_id
      ) i ON r.id = i.raw_material_id;
END;"""
)

p("Besides the database trigger, the application has several «business triggers» "
  "implemented in the service layer: the view production_create checks "
  "raw-material sufficiency before producing; sale_create forbids a sale below "
  "cost and updates the budget; salary_create deducts the salary from the "
  "budget; loan_create increases the budget; the model BusinessLoan refuses to "
  "store more than three loans simultaneously. Together they form an additional "
  "layer of business rules that guarantee data consistency.")


doc.add_page_break()


# ============================================================
# 3. CLIENT SIDE
# ============================================================
h1("3. Client side")

h2("3.1. Description of project creation")
p("The web application is developed in Python using the Django 5.2.7 framework "
  "and the classic MVT (Model-View-Template) architectural pattern, which is "
  "Django’s flavour of MVC. The project is managed with pip / venv; dependencies "
  "are listed in requirements.txt.")

p("Used tools and technologies:", bold=False, indent=False)
bullet("Python 3.13 — the programming language used for the project.")
bullet("Django 5.2.7 — a full-stack web framework that provides ORM, routing, "
       "templates, forms, authentication and the admin site out of the box.")
bullet("Django ORM — object-relational mapping that maps Python classes to MSSQL "
       "tables and generates migrations.")
bullet("mssql-django + ODBC Driver 17 for SQL Server — JDBC-style driver that "
       "lets Django talk to Microsoft SQL Server.")
bullet("Django Templates — built-in HTML templating engine used in the View "
       "layer to render the pages.")
bullet("Bootstrap 5 (CDN) — CSS framework used to style cards, tables, "
       "buttons and modals.")
bullet("django-widget-tweaks — a small package used to add CSS classes to "
       "form widgets directly in templates.")
bullet("python-dotenv — loads database credentials from a .env file.")
bullet("pip — package manager and dependency management tool (configured in "
       "requirements.txt).")
bullet("PyCharm — the IDE in which the project is developed.")

p("Steps to create the project:")
p("1. Create and activate a virtual environment:")
code_block("python -m venv venv\nvenv\\Scripts\\activate")
p("2. Install the dependencies:")
code_block("pip install django mssql-django python-dotenv")
p("3. Generate the project skeleton and the warehouse application:")
code_block("django-admin startproject clothes_factory_warehouse .\npython manage.py startapp warehouse")
p("4. Register the application in clothes_factory_warehouse/settings.py:")
code_block(
"""INSTALLED_APPS = [
    'django.contrib.admin',
    'django.contrib.auth',
    'django.contrib.contenttypes',
    'django.contrib.sessions',
    'django.contrib.messages',
    'django.contrib.staticfiles',
    'warehouse',
    'widget_tweaks',
]"""
)
p("5. Configure the MSSQL connection using credentials read from .env:")
code_block(
"""DATABASES = {
    "default": {
        "ENGINE": "mssql",
        "NAME":   os.getenv("DB_NAME"),
        "HOST":   os.getenv("DB_HOST"),
        "PORT":   os.getenv("DB_PORT"),
        "OPTIONS": {
            "driver": "ODBC Driver 17 for SQL Server",
            "extra_params":
                "Trusted_Connection=yes;TrustServerCertificate=yes",
        },
    }
}"""
)
p("6. Run the migrations: the 0001_initial migration creates all 14 domain "
  "tables and Django service tables; 0002_create_sql_logic provisions the "
  "stored procedures, the trigger and the three views; 0003_seed_workwear_data "
  "fills the catalogues with a workwear demo dataset; later migrations add "
  "loans, salary payments, production requests, the role field on Position "
  "and the UserProfile entity.")
code_block("python manage.py migrate")
p("7. Create a Django superuser to bootstrap the system:")
code_block("python manage.py createsuperuser")
p("8. Run the development server:")
code_block("python manage.py runserver 8080")

p("Resulting project structure:")
code_block(
"""WarehouseOfClothes/
├── clothes_factory_warehouse/   # project settings
│   ├── settings.py
│   ├── urls.py
│   └── wsgi.py
├── warehouse/                   # business application
│   ├── models.py                # database models
│   ├── views.py                 # controllers
│   ├── auth_views.py            # authentication views
│   ├── forms.py                 # Django forms
│   ├── permissions.py           # RBAC decorators
│   ├── context_processors.py    # role flags for templates
│   ├── urls.py                  # URL routing
│   ├── migrations/              # 0001..0007 schema + SQL logic
│   ├── templates/warehouse/     # HTML templates
│   └── static/                  # CSS / JS / images
├── manage.py
├── requirements.txt
└── .env                         # database credentials"""
)


h2("3.2. More detailed analysis")
p("Below each main directory of the warehouse application is analysed with "
  "example code.")


# 1. Models
h3("1. Models (Entity layer)")
p("Located in warehouse/models.py. Each class inherits from django.db.models.Model "
  "and corresponds to one table in the database. Django ORM automatically "
  "generates the schema from these classes and creates migrations for any "
  "change. Verbose names are written in Russian so that they appear correctly "
  "in the admin site and on form labels. Below is an example of one of the "
  "richer entities — RawMaterial:")
code_block(
"""class RawMaterial(models.Model):
    name      = models.CharField(max_length=200, verbose_name="Наименование сырья")
    unit      = models.ForeignKey(Unit, on_delete=models.PROTECT, verbose_name="Ед. изм.")
    quantity  = models.FloatField(default=0.0, verbose_name="Остаток")
    amount    = models.FloatField(default=0.0, verbose_name="Общая стоимость остатка")
    min_price = models.FloatField(default=0.0, verbose_name="Мин. цена за единицу")

    @property
    def unit_cost(self):
        \"\"\"Actual cost per unit. Falls back to min_price if stock is 0.\"\"\"
        if self.quantity > 0:
            return self.amount / self.quantity
        return self.min_price

    class Meta:
        verbose_name = "Сырьё"
        verbose_name_plural = "Сырьё"""
)
p("The model demonstrates several typical Django techniques:")
bullet("ForeignKey to Unit with on_delete=PROTECT — a unit cannot be deleted "
       "if there is a raw material that uses it.")
bullet("@property unit_cost — a computed attribute used everywhere when the "
       "actual unit cost is required (cost calculation, materials-needed "
       "report, sale-below-cost check).")
bullet("Meta.verbose_name — human-readable name shown by the admin and the "
       "templates.")

# 2. Views
h3("2. Views (Controllers)")
p("Located in warehouse/views.py and warehouse/auth_views.py. Views accept an "
  "HttpRequest, perform the business logic and return an HttpResponse. They "
  "use a few helper functions for generic CRUD (generic_list, generic_form, "
  "generic_delete) and specific functions for non-trivial operations "
  "(purchase_create, production_create, sale_create, salary_create, "
  "loan_create, loan_payoff, production_request_create, analytics_dashboard). "
  "Below is the production view, which checks raw-material availability "
  "before committing the operation:")
code_block(
"""@role_required(Position.ROLE_PRODUCTION)
def production_create(request):
    if request.method == 'POST':
        form = ProductionForm(request.POST)
        if form.is_valid():
            emp = _resolve_employee(request)
            production = form.save(commit=False)
            production.employee = emp
            product = production.product
            qty = production.quantity

            ingredients = product.ingredient_set.all()
            missing = []
            for ing in ingredients:
                needed = ing.quantity * qty
                if ing.raw_material.quantity < needed:
                    missing.append(
                        f"<b>{ing.raw_material.name}</b> &mdash; "
                        f"need {needed:,.2f}, available {ing.raw_material.quantity:,.2f}"
                    )

            if missing:
                messages.error(request, mark_safe(
                    "Insufficient raw materials:<ul><li>"
                    + "</li><li>".join(missing) + "</li></ul>"))
            else:
                for ing in ingredients:
                    raw = ing.raw_material
                    needed = ing.quantity * qty
                    unit_cost = raw.amount / raw.quantity
                    raw.quantity -= needed
                    raw.amount   -= needed * unit_cost
                    raw.save()

                product.quantity += qty
                product.save()
                production.save()
                messages.success(request, f"Produced {qty} of {product.name}.")
                return redirect('production_list')
    else:
        form = ProductionForm()
    return render(request, 'warehouse/form.html',
                  {'form': form, 'title': 'Новое производство'})"""
)

# 3. Forms
h3("3. Forms")
p("Located in warehouse/forms.py. Django forms generate the HTML inputs, perform "
  "server-side validation and convert POST data into clean Python types. Most "
  "forms are ModelForms — bound directly to a model — and one is a plain Form "
  "(RawPurchaseForm), because the actual database write is done by the stored "
  "procedure rather than by ORM:")
code_block(
"""class RawPurchaseForm(forms.Form):
    raw_material = forms.ModelChoiceField(
        queryset=RawMaterial.objects.all(), label="Сырьё")
    quantity = forms.FloatField(label="Количество")
    amount   = forms.FloatField(label="Сумма закупки")

class ProductSaleForm(forms.ModelForm):
    class Meta:
        model = ProductSale
        fields = ['product', 'quantity']  # amount is computed in the view

class ProductionForm(forms.ModelForm):
    class Meta:
        model = Production
        exclude = ['date', 'employee']    # date is auto, employee from request"""
)

# 4. URL routing
h3("4. URL routing")
p("Located in warehouse/urls.py. Maps each URL to a view function. CRUD URLs "
  "follow Django’s naming convention list / add / <pk>/edit / <pk>/delete and "
  "are organised by domain section. Below is a representative excerpt:")
code_block(
"""from django.urls import path
from . import views, auth_views

urlpatterns = [
    path('', views.index, name='index'),

    # Auth
    path('login/',  auth_views.login_view,  name='login'),
    path('logout/', auth_views.logout_view, name='logout'),

    # Reference data
    path('units/',       views.unit_list,        name='unit_list'),
    path('positions/',   views.position_list,    name='position_list'),
    path('employees/',   views.employee_list,    name='employee_list'),

    # Stock
    path('raw_materials/',     views.raw_material_list,     name='raw_material_list'),
    path('finished_products/', views.finished_product_list, name='finished_product_list'),
    path('ingredients/',       views.ingredient_list,       name='ingredient_list'),

    # Operations
    path('purchase/add/',  views.purchase_create,  name='purchase_create'),
    path('production/',    views.production_list,  name='production_list'),
    path('sale/add/',      views.sale_create,      name='sale_create'),
    path('salaries/',      views.salary_list,      name='salary_list'),
    path('loans/',         views.loan_list,        name='loan_list'),

    # SQL Views (raw SQL reports)
    path('view/raw_stock/',        views.view_raw_stock,        name='view_raw_stock'),
    path('view/product_stock/',    views.view_product_stock,    name='view_product_stock'),
    path('view/materials_needed/', views.view_materials_needed, name='view_materials_needed'),

    # Analytics dashboard
    path('analytics/dashboard/', views.analytics_dashboard, name='analytics_dashboard'),
]"""
)

# 5. Permissions
h3("5. Permissions and roles (RBAC)")
p("Located in warehouse/permissions.py. Every Position has a role attribute "
  "(director, accountant, warehouse, production, sales, none). Each Django "
  "auth_user can be linked to one Employee via UserProfile, from which the "
  "system resolves the user’s role. Two decorators wrap the views to enforce "
  "access:")
code_block(
"""def role_required(*allowed_roles):
    \"\"\"Decorator: only users with one of the roles (director always passes).\"\"\"
    def decorator(view_func):
        @wraps(view_func)
        @login_required
        def _wrapped(request, *args, **kwargs):
            if has_role(request.user, *allowed_roles):
                return view_func(request, *args, **kwargs)
            messages.error(request, "У вас нет прав на это действие.")
            return redirect('index')
        return _wrapped
    return decorator

def director_required(view_func):
    \"\"\"Decorator: only director / superuser.\"\"\"
    @wraps(view_func)
    @login_required
    def _wrapped(request, *args, **kwargs):
        if is_director(request.user):
            return view_func(request, *args, **kwargs)
        messages.error(request, "Только для директора.")
        return redirect('index')
    return _wrapped"""
)
p("Templates also adapt to the role thanks to context_processors.auth_context, "
  "which exposes flags is_director, can_warehouse, can_production, can_sales, "
  "can_accountant for every page — this is how irrelevant menu items disappear "
  "for users with restricted roles.")

# 6. Templates
h3("6. Templates (Presentation)")
p("Located in warehouse/templates/. The shared layout warehouse/templates/"
  "base.html includes a top navigation bar that is built from the role flags "
  "above. Each domain page extends base.html and provides its content block. "
  "Below is a fragment of the dashboard (warehouse/index.html):")
code_block(
"""{% if is_director %}
<div class="col-md-4">
  <div class="card h-100 border-primary">
    <div class="card-header bg-primary text-white">📘 Справочники</div>
    <div class="list-group list-group-flush">
      <a href="{% url 'unit_list' %}"    class="list-group-item">📏 Единицы измерения</a>
      <a href="{% url 'position_list' %}" class="list-group-item">👔 Должности</a>
      <a href="{% url 'employee_list' %}" class="list-group-item">👥 Сотрудники</a>
      <a href="{% url 'account_list' %}"  class="list-group-item">👤 Аккаунты</a>
    </div>
  </div>
</div>
{% endif %}

{% if is_director or can_warehouse %}
<a href="{% url 'raw_material_list' %}"
   class="list-group-item">🧶 Сырьё</a>
{% endif %}"""
)
p("Templates are rendered through Django’s template engine. Bootstrap 5 is "
  "loaded from the CDN, plus a custom CSS block in base.html that defines a "
  "soft-gradient palette and elevated cards. Forms are rendered with django-"
  "widget-tweaks (|add_class) to attach Bootstrap classes without rewriting "
  "the form widgets in Python.")


doc.add_page_break()


# ============================================================
# 4. LAUNCHING THE PROGRAM
# ============================================================
h1("4. Launching the program")
p("To launch the project on localhost the following steps are required:")
p("1. Make sure the MSSQL instance is running and a database matching DB_NAME "
   "in the .env file exists. Driver «ODBC Driver 17 for SQL Server» must be "
   "installed.")
p("2. Activate the virtual environment and install the dependencies:")
code_block("venv\\Scripts\\activate\npip install -r requirements.txt")
p("3. Run the migrations — this creates the tables, the stored procedures, "
   "the trigger, the views and seeds the catalogues:")
code_block("python manage.py migrate")
p("4. Create the first director (superuser) so that an administrator can log in:")
code_block("python manage.py createsuperuser")
p("5. Start the development server:")
code_block("python manage.py runserver 8080")
p("6. Open http://127.0.0.1:8080 in the browser, log in with the superuser "
   "account, then via Аккаунты create accounts for the rest of the staff.")


doc.add_page_break()


# ============================================================
# 5. SCREENSHOTS
# ============================================================
h1("5. Screenshots of the program")

p("Below the main pages of the application are presented. Screenshots are "
  "captured against the workwear demo dataset that is seeded by migration "
  "0003_seed_workwear_data.")

screens = [
    ("Главная страница (Дашборд).png",
     "Image 19. Dashboard — the main page seen after a successful login. "
     "Tiles and menu items are shown only for the sections that the user’s "
     "role grants access to."),
    ("Состояние бюджета.png",
     "Image 20. Budget page — read and edit the single budget record. Access "
     "is restricted to the accountant and the director."),
    ("Справочник должностей.png",
     "Image 21. Positions reference — each position is bound to a role that "
     "drives RBAC."),
    ("Справочник сотрудников.png",
     "Image 22. Employees reference — full information about staff members."),
    ("Справочник единиц измерения.png",
     "Image 23. Units of measurement reference."),
    ("Справочник сырья.png",
     "Image 24. Raw materials reference — current stock, total stock value "
     "and minimum (fallback) unit price."),
    ("Справочник готовой продукции.png",
     "Image 25. Finished products reference."),
    ("Состав продукции (Ингредиенты).png",
     "Image 26. Recipes — how much of each raw material is required to "
     "produce one unit of a finished product."),
    ("Закупка сырья (Форма).png",
     "Image 27. Raw material purchase form — the actual INSERT goes through "
     "the sp_create_purchase stored procedure."),
    ("Производство.png",
     "Image 28. Production page — manufacturing finished products from raw "
     "materials. The view rejects the operation if any of the ingredients is "
     "not enough."),
    ("Оформление продажи готовой продукции.png",
     "Image 29. Sale page — sells finished products. The price is automatically "
     "computed as unit_cost × 1.30 × qty; the budget is increased."),
    ("Выплата зарплаты.png",
     "Image 30. Salary payouts — each payout deducts the corresponding amount "
     "from the budget."),
    ("Бизнес-кредиты.png",
     "Image 31. Business loans — taking a new loan increases the budget; "
     "paying it off (button «Погасить») decreases the budget."),
    ("Аналитический дашборд.png",
     "Image 32. Analytics dashboard — chart of top-sold products and the "
     "ratio of revenue to expenses, plus drill-down tables of sales, "
     "purchases and salary payouts."),
]
for filename, cap in screens:
    picture(os.path.join(SHOTS, filename), 6.0, cap)


doc.add_page_break()


# ============================================================
# 6. CONCLUSION
# ============================================================
h1("6. Conclusion")
p("During the development of this web application as part of the term project, "
  "valuable experience was gained in working with the Django framework, the "
  "Python programming language and the Microsoft SQL Server DBMS. Skills in "
  "designing relational schemas, writing stored procedures, triggers and "
  "views, and connecting them to an application were significantly improved.")
p("The server side was implemented on MSSQL using SQL Server Management Studio "
  "and Django migrations (RunSQL), while the client side was built with Python "
  "and Django (MVT pattern) in PyCharm. Domain logic was distributed across "
  "the layers in the following way: data integrity rules live in the DBMS "
  "(stored procedures, trigger, FK constraints, views), while business workflow "
  "and authorisation live in Django (RBAC decorators, model save guards, "
  "service-layer code in the views).")
p("The result is a self-contained system that covers the full workflow of a "
  "garment factory: from raw-material purchase to finished-product sale, with "
  "real-time budget control, role-based access and analytics. The architecture "
  "follows clean-architecture principles — models, forms, views, permissions "
  "and templates are organised in separate modules.")


# ============================================================
# 7. LITERATURE
# ============================================================
h1("7. List of Used Literature")
p("1. Django Software Foundation. Django 5.2 documentation. — "
  "https://docs.djangoproject.com/en/5.2/", indent=False)
p("2. Microsoft. Transact-SQL Reference (CREATE PROCEDURE, CREATE TRIGGER, "
  "CREATE VIEW). — https://learn.microsoft.com/en-us/sql/t-sql/", indent=False)
p("3. Microsoft. mssql-django driver documentation. — "
  "https://github.com/microsoft/mssql-django", indent=False)
p("4. Bootstrap. Bootstrap 5 documentation. — https://getbootstrap.com/docs/5.3/",
  indent=False)
p("5. Mark Lutz. Learning Python, 5th ed. — O’Reilly Media, 2013.", indent=False)
p("6. Itzik Ben-Gan. T-SQL Fundamentals, 4th ed. — Microsoft Press, 2020.", indent=False)
p("7. Training materials of the department «Software of computer systems», "
  "KSTU named after I. Razzakov.", indent=False)


# ============================================================
# 8. APPLICATION APPENDIX
# ============================================================
doc.add_page_break()
h1("8. Application appendix")
p("This appendix lists additional source files that are useful for a more "
  "complete acquaintance with the project.")

h2("1. Production-request workflow (process_production_request)")
p("The function below is the heart of the multi-step production request "
  "feature: it walks the request through five statuses (checking raw stock → "
  "purchasing → producing → selling → completed), automatically purchases "
  "missing raw materials if the budget allows it, performs the production and "
  "the sale, and updates the budget at every step.")
code_block(
"""def process_production_request(req, acting_employee=None):
    try:
        req.status = ProductionRequest.STATUS_CHECKING_RAW
        req.save()

        product = req.product
        qty = req.quantity
        ingredients = product.ingredient_set.select_related('raw_material').all()

        # 1. Checking raw stock
        missing_materials = []
        total_cost_of_production = 0
        for ing in ingredients:
            needed = ing.quantity * qty
            available = ing.raw_material.quantity
            unit_cost = ing.raw_material.unit_cost
            total_cost_of_production += needed * unit_cost
            if needed > available:
                missing_materials.append({
                    'raw_material': ing.raw_material,
                    'qty':  needed - available,
                    'cost': (needed - available) * unit_cost,
                })

        # 2. Purchasing missing materials (if budget allows)
        if missing_materials:
            req.status = ProductionRequest.STATUS_PURCHASING
            req.save()
            total_purchase_cost = sum(m['cost'] for m in missing_materials)
            budget = Budget.objects.first()
            if not budget or budget.amount < total_purchase_cost:
                req.rejection_reason = "Не хватило бюджета на закупку сырья."
                req.save()
                return
            budget.amount -= total_purchase_cost
            budget.save()
            employee = acting_employee or Employee.objects.first()
            for m in missing_materials:
                raw = m['raw_material']
                RawPurchase.objects.create(
                    raw_material=raw, quantity=m['qty'],
                    amount=m['cost'], employee=employee)
                raw.quantity += m['qty']
                raw.amount   += m['cost']
                raw.save()

        # 3. Production
        req.status = ProductionRequest.STATUS_PRODUCTION
        req.save()
        for ing in ingredients:
            raw = ing.raw_material
            needed = ing.quantity * qty
            unit_cost = raw.unit_cost
            raw.quantity -= needed
            raw.amount   -= needed * unit_cost
            raw.save()
        Production.objects.create(
            product=product, quantity=qty,
            employee=Employee.objects.first())
        product.quantity += qty
        product.amount   += total_cost_of_production
        product.save()

        # 4. Sale
        req.status = ProductionRequest.STATUS_SALE
        req.save()
        sale_price = total_cost_of_production * 1.30
        budget = Budget.objects.first()
        budget.amount += sale_price
        budget.save()
        ProductSale.objects.create(
            product=product, quantity=qty,
            amount=sale_price, employee=Employee.objects.first())
        product.quantity -= qty
        product.amount   -= total_cost_of_production
        product.save()

        # 5. Done
        req.status = ProductionRequest.STATUS_COMPLETED
        req.save()
    except Exception as e:
        req.status = ProductionRequest.STATUS_ERROR
        req.rejection_reason = f"Системный сбой: {e}"
        req.save()"""
)

h2("2. Analytics dashboard view")
p("The view aggregates sales, purchases and salaries and serialises the data "
  "for the Chart.js charts on the page.")
code_block(
"""@role_required(Position.ROLE_ACCOUNTANT)
def analytics_dashboard(request):
    sales     = ProductSale.objects.select_related('product', 'employee').order_by('-date')
    purchases = RawPurchase.objects.select_related('raw_material', 'employee').order_by('-date')
    salaries  = SalaryPayment.objects.select_related('employee').order_by('-date')

    # Top-5 sold products
    top_products_qs = (sales
        .values('product__name')
        .annotate(total_qty=Sum('quantity'), total_revenue=Sum('amount'))
        .order_by('-total_revenue')[:5])
    top_labels = [x['product__name']     for x in top_products_qs]
    top_data   = [float(x['total_revenue']) for x in top_products_qs]

    total_sales     = sales.aggregate(Sum('amount'))['amount__sum']     or 0
    total_purchases = purchases.aggregate(Sum('amount'))['amount__sum'] or 0
    total_salaries  = salaries.aggregate(Sum('amount'))['amount__sum']  or 0

    return render(request, 'warehouse/analytics_dashboard.html', {
        'sales': sales, 'purchases': purchases, 'salaries': salaries,
        'chart_top_products_labels': json.dumps(top_labels, cls=DjangoJSONEncoder),
        'chart_top_products_data':   json.dumps(top_data,   cls=DjangoJSONEncoder),
        'total_sales':     float(total_sales),
        'total_purchases': float(total_purchases),
        'total_salaries':  float(total_salaries),
        'total_expenses':  float(total_purchases + total_salaries),
    })"""
)

h2("3. Base layout (base.html, abridged)")
p("The shared layout provides the top navigation bar, the message strip and "
  "the role-driven menu items. Bootstrap 5 is loaded from the CDN; a custom "
  "<style> block defines the colour palette and the elevated-card design.")
code_block(
"""<nav class="navbar navbar-expand-lg navbar-dark">
  <div class="container">
    <a class="navbar-brand" href="{% url 'index' %}">Warehouse Of Clothes</a>
    <div class="collapse navbar-collapse">
      <ul class="navbar-nav ms-auto">
        {% if user.is_authenticated %}
          {% if can_accountant %}
            <li class="nav-item"><a class="nav-link" href="{% url 'budget_list' %}">💰 Бюджет</a></li>
            <li class="nav-item"><a class="nav-link" href="{% url 'salary_list' %}">💸 Зарплаты</a></li>
            <li class="nav-item"><a class="nav-link" href="{% url 'loan_list' %}">🏦 Кредиты</a></li>
          {% endif %}
          {% if is_director %}
            <li class="nav-item"><a class="nav-link" href="{% url 'account_list' %}">👤 Аккаунты</a></li>
          {% endif %}
          <li class="nav-item"><span class="badge bg-success">
            {{ current_employee.full_name|default:user.username }}</span></li>
          <li class="nav-item"><a class="nav-link" href="{% url 'logout' %}">🚪 Выйти</a></li>
        {% else %}
          <li class="nav-item"><a class="nav-link" href="{% url 'login' %}">🔐 Войти</a></li>
        {% endif %}
      </ul>
    </div>
  </div>
</nav>

<div class="container mt-5">
    {% if messages %}{% for message in messages %}
        <div class="alert alert-{{ message.tags }}">{{ message }}</div>
    {% endfor %}{% endif %}
    {% block content %}{% endblock %}
</div>"""
)


# save ---------------------------------------------------------
out = os.path.join(BASE, "DBMS_TermProject_Klyuchevsky_WarehouseOfClothes.docx")
doc.save(out)
print(f"Saved: {out}")
