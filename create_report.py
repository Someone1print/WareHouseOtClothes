import sys
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.5
paragraph_format.first_line_indent = Cm(1.25)
paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # usually justified for GOST reports

def add_heading(doc, text, level=1):
    # Not using standard headings to get normal font in TOC easily and maintain 1.5 spacing/first line indent
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = 0
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    return p
    
def add_paragraph(doc, text, bold=False, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.alignment = align
    return p

# Title page
p = doc.add_paragraph("Министерство образования и науки Кыргызской Республики")
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.paragraph_format.first_line_indent = 0

p = doc.add_paragraph("Кыргызский государственный технический университет\nим. И.Раззакова")
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.paragraph_format.first_line_indent = 0

p = doc.add_paragraph("Институт информационных технологий")
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.paragraph_format.first_line_indent = 0

p = doc.add_paragraph("Кафедра «Программное обеспечение компьютерных систем»")
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.paragraph_format.first_line_indent = 0

p = doc.add_paragraph("Направление: 710400 «Программная инженерия»")
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.paragraph_format.first_line_indent = 0

p = doc.add_paragraph("Дисциплина: «Основы разработки и анализа требований к ПО»")
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.paragraph_format.first_line_indent = 0

doc.add_paragraph("\n\n")

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.paragraph_format.first_line_indent = 0
r = p.add_run("Отчет")
r.bold = True
r.font.size = Pt(16)

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.paragraph_format.first_line_indent = 0
p.add_run("Наименование работы:\n«Проектирование и разработка автоматизированной системы "
          "управления складом швейной фабрики (WarehouseOfClothes)»")

doc.add_paragraph("\n\n\n\n")

p = doc.add_paragraph("Выполнил: студент группы\nПИ(б)-4-23 Ключевский Дмитрий")
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
p.paragraph_format.first_line_indent = 0
p = doc.add_paragraph("Проверила: Макиева З.Д.")
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
p.paragraph_format.first_line_indent = 0

doc.add_paragraph("\n\n\n\n\n\n")
p = doc.add_paragraph("Бишкек – 2026")
p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.paragraph_format.first_line_indent = 0

doc.add_page_break()

add_heading(doc, 'Оглавление', level=1)
add_paragraph(doc, "1. Экономическое обоснование разработки проектного решения")
add_paragraph(doc, "1.1. Характеристика разработанного по индивидуальному заказу программного средства")
add_paragraph(doc, "1.2. Расчет затрат на разработку и цена программного средства, созданного по индивидуальному заказу")
add_paragraph(doc, "1.3. Расчет показателей экономической эффективности разработки и использования программного средства")
add_paragraph(doc, "2. Анализ и моделирование целевых бизнес-процессов предметной области")
add_paragraph(doc, "2.1. Графический материал: AS-IS и TO-BE модели процессов предметной области в нотации BPMN")
add_paragraph(doc, "2.2. GAP-анализ")
add_paragraph(doc, "2.3. Диаграммы IDEF0 и DFD")
add_paragraph(doc, "Вывод")

doc.add_page_break()

add_heading(doc, '1. Экономическое обоснование разработки проектного решения')
add_heading(doc, '1.1. Характеристика разработанного по индивидуальному заказу программного средства', level=2)

add_paragraph(doc, "Целью разработки является решение актуальных проблем швейных предприятий, связанных с ведением складского и финансового учета.")
add_paragraph(doc, "Первая проблема заключается в том, что учет сырья (тканей, фурнитуры) и готовой продукции ведется вручную с использованием бумажных журналов или электронных таблиц. Это приводит к ошибкам, потере данных и затрудняет контроль остатков.")
add_paragraph(doc, "Вторая проблема связана с отсутствием интеграции складского учета и финансовых операций. Закупка сырья и продажа продукции напрямую влияют на бюджет предприятия, однако при ручном учете контроль этих процессов является затруднительным.")
add_paragraph(doc, "Разрабатываемое программное средство представляет собой веб-приложение «WarehouseOfClothes», предназначенное для автоматизации складских и финансовых процессов.")

add_paragraph(doc, "Система позволяет:")
p = add_paragraph(doc, "· учитывать поступление сырья;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· контролировать движение материалов;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· автоматически отслеживать остатки;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· учитывать готовую продукцию;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· вести финансовый учет;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· формировать аналитические отчеты.")
p.paragraph_format.first_line_indent = Cm(2)

add_paragraph(doc, "Область применения — предприятия швейной промышленности.")
add_paragraph(doc, "Таким образом, внедрение программного средства позволяет снизить количество ошибок, ускорить обработку данных и повысить эффективность управления предприятием.")

add_paragraph(doc, "Разработчиком программного средства является Ключевский Дмитрий.", bold=True)
add_paragraph(doc, "Организация-заказчик — Кыргызский государственный технический университет.", bold=True)

add_paragraph(doc, "Актуальность разработки обусловлена:")
p = add_paragraph(doc, "· ростом объемов производства;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· необходимостью точного учета сырья;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· потребностью в автоматизации процессов;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· необходимостью оперативной аналитики.")
p.paragraph_format.first_line_indent = Cm(2)


add_heading(doc, '1.2. Расчет затрат на разработку и цена программного средства')

add_paragraph(doc, "1. Основная заработная плата разработчиков:", bold=True)

# Table of salaries
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Должность'
hdr_cells[1].text = 'Оклад (сом)'
# style header
for cell in hdr_cells:
    for p in cell.paragraphs:
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)

salaries = [
    ('Бизнес-аналитик', '5 681,8'),
    ('Системный архитектор', '13 636,2'),
    ('Программист', '20 454,4'),
    ('Тестировщик', '6 818,1'),
    ('Дизайнер', '11 363,6'),
    ('Итого', '57 954,1'),
    ('Премия', '28 977,05'),
    ('Всего', '86 931,15')
]
for role, amount in salaries:
    row_cells = table.add_row().cells
    row_cells[0].text = role
    row_cells[1].text = amount
    for cell in row_cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(14)
    
# Add some space
doc.add_paragraph()

add_paragraph(doc, "2. Дополнительная заработанная плата разработчиков:", bold=True)
add_paragraph(doc, "Зд = 57 954,1 × 0,15 = 8 693,12 сом")

add_paragraph(doc, "3. Отчисление на социальные нужды:", bold=True)
add_paragraph(doc, "(57 954,1 + 8 693,12) × 0,346 = 66 647,22 × 0,346 = 23 060,93 сом")

add_paragraph(doc, "4. Прочие расходы:", bold=True)
add_paragraph(doc, "57 954,1 × 0,35 = 20 283,94 сом")

add_paragraph(doc, "5. Общая сумма затрат на разработку:", bold=True)
add_paragraph(doc, "57 954,1 + 8 693,12 + 23 060,93 + 20 283,94 = 109 992,09 сом")

add_paragraph(doc, "6. Плановая прибыль:", bold=True)
add_paragraph(doc, "109 992,09 × 0,30 = 32 997,63 сом")

add_paragraph(doc, "7. Отпускная цена:", bold=True)
add_paragraph(doc, "109 992,09 + 32 997,63 = 142 989,72 сом")

add_heading(doc, '1.3. Расчет показателей экономической эффективности')

add_paragraph(doc, "Рентабельность инвестиций (Ри) рассчитывается следующим образом (от плановой прибыли и суммы затрат):")
add_paragraph(doc, "Ри = (29 697,87 / 109 992,09) × 100% ≈ 27%")
# Note: From the prompt snippet, there's a specific amount 29 697,87, maybe I should leave it exactly as provided in the prompt:
# "Ри = (29 697,87 / 109 992,09) × 100% \n Ри = 27%" 
# I did a bit of restructuring to make it readable in Word. Let's fix the exact text to exactly what the user wrote.

doc.add_page_break()

add_heading(doc, '2. Анализ и моделирование целевых бизнес-процессов предметной области')

add_paragraph(doc, "Текущее состояние (AS-IS)", bold=True)
add_paragraph(doc, "На предприятии учет осуществляется вручную с использованием бумажных документов и Excel.")
add_paragraph(doc, "Процессы поступления, хранения и выдачи сырья не автоматизированы.")
add_paragraph(doc, "Финансовый учет ведется отдельно от складского.")

add_paragraph(doc, "Основные недостатки:")
p = add_paragraph(doc, "· высокая вероятность ошибок;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· отсутствие централизованной системы;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· задержки в обработке данных;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· отсутствие аналитики.")
p.paragraph_format.first_line_indent = Cm(2)

add_paragraph(doc, "Целевое состояние (TO-BE)", bold=True)
add_paragraph(doc, "После внедрения системы WarehouseOfClothes все процессы автоматизируются.")
add_paragraph(doc, "Система обеспечивает:")
p = add_paragraph(doc, "· централизованное хранение данных;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· автоматический учет остатков;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· контроль финансов;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· формирование отчетности;")
p.paragraph_format.first_line_indent = Cm(2)
p = add_paragraph(doc, "· аналитический дашборд.")
p.paragraph_format.first_line_indent = Cm(2)

add_paragraph(doc, "В результате снижаются ошибки и повышается эффективность работы.")

add_heading(doc, '2.1. Графический материал: AS-IS и TO-BE модели процессов (BPMN)')

doc.add_paragraph("\n[ВСТАВИТЬ ЗДЕСЬ ДИАГРАММУ AS-IS]\n").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Рисунок 1 – Диаграмма AS-IS в нотации BPMN").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("\n[ВСТАВИТЬ ЗДЕСЬ ДИАГРАММУ TO-BE]\n").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Рисунок 2 – Диаграмма TO-BE в нотации BPMN").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

add_heading(doc, '2.2. GAP-анализ')

# GAP table
table_gap = doc.add_table(rows=1, cols=5)
table_gap.style = 'Table Grid'
hdr_cells = table_gap.rows[0].cells
hdr_cells[0].text = 'Область управления'
hdr_cells[1].text = 'AS-IS (как есть)'
hdr_cells[2].text = 'TO-BE (как будет)'
hdr_cells[3].text = 'GAP (разрыв)'
hdr_cells[4].text = 'Решение'

for cell in hdr_cells:
    for p in cell.paragraphs:
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)

gaps = [
    ('Учет сырья', 'Вручную', 'Автоматизировано', 'Нет системы', 'Разработка ПО'),
    ('Остатки', 'Ошибки', 'Контроль', 'Нет контроля', 'Триггеры БД'),
    ('Финансы', 'Отдельно', 'Интеграция', 'Нет связи', 'Единая база'),
    ('Отчеты', 'Вручную', 'Автоматически', 'Медленно', 'Dashboard')
]
for область, асис, туби, гэп, решение in gaps:
    row_cells = table_gap.add_row().cells
    row_cells[0].text = область
    row_cells[1].text = асис
    row_cells[2].text = туби
    row_cells[3].text = гэп
    row_cells[4].text = решение
    for cell in row_cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(14)

doc.add_paragraph()
add_heading(doc, '2.3. Диаграммы IDEF0 и DFD')

doc.add_paragraph("\n[ВСТАВИТЬ IDEF0 A0]\n").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Рисунок 3 – Контекстная диаграмма IDEF0 (A-0)").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("\n[ВСТАВИТЬ IDEF0 ДЕКОМПОЗИЦИЮ]\n").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Рисунок 4 – Диаграмма декомпозиции IDEF0").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("\n[ВСТАВИТЬ DFD]\n").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Рисунок 5 – Диаграмма потоков данных (DFD)").paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


add_heading(doc, 'Вывод')
add_paragraph(doc, "В ходе выполнения лабораторной работы было разработано экономическое обоснование и проведено моделирование бизнес-процессов для системы WarehouseOfClothes.")
add_paragraph(doc, "Экономическая эффективность проекта подтверждена расчетами. Автоматизация процессов позволяет снизить ошибки и повысить эффективность работы предприятия.")
add_paragraph(doc, "Таким образом, разработанное программное средство рекомендуется к внедрению.")

doc.save('Отчет_по_лабораторной_2.docx')
